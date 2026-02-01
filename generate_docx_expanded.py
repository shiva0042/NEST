
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

def add_code_files(doc, root_dir):
    """
    Traverses the directory and adds code files to the document.
    """
    doc.add_section()
    header = doc.add_heading('APPENDIX A: SOURCE CODE', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Walk through the directory
    for dirpath, dirnames, filenames in os.walk(root_dir):
        # Sort for consistent order
        filenames.sort()
        for filename in filenames:
            if filename.endswith('.dart'):
                file_path = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(file_path, root_dir)
                
                # Add Filename Header
                p = doc.add_paragraph()
                p.add_run(f"File: {rel_path}").bold = True
                p.style = 'Heading 3'
                
                # Add Code Content
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        code_content = f.read()
                    
                    code_paragraph = doc.add_paragraph(code_content)
                    code_paragraph.style = 'No Spacing' 
                    code_font = code_paragraph.runs[0].font
                    code_font.name = 'Courier New'
                    code_font.size = Pt(8) # Small font for code to fit more horizontally
                    
                except Exception as e:
                    doc.add_paragraph(f"[Error reading file: {e}]")
                
                doc.add_page_break()

def create_report():
    # Read the markdown file
    md_file_path = 'PROJECT_REPORT_FULL.md'
    if not os.path.exists(md_file_path):
        print(f"Error: {md_file_path} not found.")
        return

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # --- STYLE SETUP ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- TITLE PAGE ---
    # Simplified title page logic (reused)
    doc.add_section()
    for _ in range(5): doc.add_paragraph() # Spacing
    
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("PROJECT REPORT\n\nON\n\n")
    title_run.font.size = Pt(16)
    
    title_run_2 = title_paragraph.add_run("NEST (Near Easy Shop Tracker)\n")
    title_run_2.font.size = Pt(24)
    title_run_2.font.bold = True
    
    doc.add_paragraph("\n" * 5)
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Submitted in partial fulfillment of the requirements for the degree of\n\n")
    info.add_run("BACHELOR OF TECHNOLOGY\n").bold = True
    info.add_run("IN\nCOMPUTER SCIENCE AND ENGINEERING\n\n")
    
    doc.add_page_break()

    # --- MAIN CONTENT FROM MARKDOWN ---
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Main Title in Markdown -> Title Page already done, maybe ignore or treat as H1
            pass 
        elif line.startswith('## '):
            # H1 -> Start new page for major chapters
            doc.add_page_break()
            h = doc.add_heading(line[3:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or (len(line)>2 and line[1]=='.'):
             # Simple heuristic for numbered lists
            p = doc.add_paragraph(line.split('.', 1)[1].strip(), style='List Number')
        else:
            # Paragraph text with basic bolding
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: 
                    run.font.bold = True
                
    # --- APPEND SOURCE CODE ---
    # This is the key to reaching 100 pages
    add_code_files(doc, os.path.join(os.getcwd(), 'lib'))

    # Save
    output_path = 'NEST_Project_Report_100Pages.docx'
    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == "__main__":
    create_report()
