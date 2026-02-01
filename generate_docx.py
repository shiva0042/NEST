
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def create_report():
    # Read the markdown file
    md_file_path = 'PROJECT_REPORT_FULL.md'
    if not os.path.exists(md_file_path):
        print(f"Error: {md_file_path} not found.")
        return

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # --- TITLE PAGE ---
    doc.add_section()
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("PROJECT REPORT\n\n")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0) # Black

    subtitle_run = title_paragraph.add_run("NEST (Near Easy Shop Tracker)\n")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    
    # Space
    doc.add_paragraph("\n" * 5)
    
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run("Submitted by:\n[Your Name/Team]\n\nDate:\n[Current Date]")
    info_run.font.size = Pt(14)
    
    doc.add_page_break()

    # --- CONTENT ---
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Title or Main Header (already handled title page, but treating as Heading 1)
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('* ') or line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            # Numbered list logic (simplified)
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            # Normal text
            # quick bold handling for **text**
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are inside ** so bold them
                    run.font.bold = True

    # Save
    output_path = 'NEST_Project_Report.docx'
    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == "__main__":
    create_report()
